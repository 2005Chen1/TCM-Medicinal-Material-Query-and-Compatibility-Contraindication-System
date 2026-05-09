"""
生成需求分析.docx — 使用 python-docx 创建课程设计报告 Word 文档。
包含：需求分析、面向对象设计、软件体系结构风格与设计模式。
4+1视图部分见 plantuml/*.txt 文件。
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 全局样式设置 ─────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

# 标题样式
for level, (size, bold) in enumerate([(22, True), (16, True), (14, True), (12, True)], 1):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = '黑体'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading_style.font.size = Pt(size)
    heading_style.font.bold = bold
    heading_style.font.color.rgb = RGBColor(0, 0, 0)

# 辅助函数
def add_para(text, bold=False, indent=False, font_size=None, alignment=None):
    """添加段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.font.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if alignment is not None:
        p.alignment = alignment
    return p

def add_code_block(code_text):
    """添加代码块（灰底等宽字体段落）。"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        # 灰色底纹
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        shading.set(qn('w:val'), 'clear')
        p.paragraph_format.element.get_or_add_pPr().append(shading)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

def add_table(headers, rows):
    """添加表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('中医药常用药材查询与配伍禁忌系统')
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(26)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('系统需求分析与设计方案')
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(18)

for _ in range(3):
    doc.add_paragraph()

info_lines = [
    '课程名称：软件工程综合实践',
    '所属院校：南京中医药大学',
    '开发语言：Python（纯标准库，零第三方依赖）',
    '技术栈：SQLite3 + Tkinter + JSON',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 目录占位
# ═══════════════════════════════════════════════════════════════
doc.add_heading('目录', level=1)
add_para('（请在 Word 中插入自动目录：引用 → 目录 → 自动目录）')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第一章：需求分析
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第一章  需求分析', level=1)

# 1.1 项目背景
doc.add_heading('1.1 项目背景', level=2)
add_para(
    '中医药是中华民族的瑰宝，具有数千年的临床实践历史。在中药临床应用中，药材之间的配伍禁忌——'
    '尤其是"十八反""十九畏"等经典禁忌规则——是保障用药安全的关键环节。传统上，这些知识依赖医师'
    '的记忆与经验判断，缺乏自动化、信息化的辅助工具，存在以下痛点：',
    indent=True
)
add_para('（1）药材种类繁多，功效、性味、归经等信息量大，人工记忆容易遗漏或混淆。')
add_para('（2）配伍禁忌规则分散在古籍与教材中，缺乏集中查询的手段。')
add_para('（3）开具复方时，逐对检查药材冲突效率低下，容易因疏忽导致配伍风险。')
add_para(
    '针对上述问题，本系统旨在开发一个桌面端的中医药药材信息管理与配伍禁忌自动检查平台。系统以'
    'Python标准库实现，无需安装任何第三方依赖，适用于中医药从业人员、学生及爱好者在本地计算机上'
    '查阅药材信息、管理禁忌规则、快速识别处方中的配伍风险。',
    indent=True
)

# 1.2 功能需求
doc.add_heading('1.2 功能需求', level=2)
add_para('系统共规划六大核心功能模块，以用例描述方式逐一说明如下。')

# 模块一
doc.add_heading('1.2.1 模块一：药材管理', level=3)
add_table(
    ['项目', '内容'],
    [
        ['用例名称', '药材信息增删改查'],
        ['参与者', '系统用户（中医药从业人员、学生、爱好者）'],
        ['前置条件', '数据库已初始化，系统已启动'],
        ['主流程',
         '1. 用户打开"药材管理"标签页\n'
         '2. 系统展示全部药材列表（ID / 名称 / 拼音 / 分类 / 功效）\n'
         '3. 用户选择列表中某药材，右侧表单自动填充全部属性（10个字段）\n'
         '4. 用户可修改表单后点击"更新所选"保存，或点击"删除所选"移除\n'
         '5. 用户也可清空表单填写新药材，点击"新增药材"添加'],
        ['后置条件', '药材列表实时刷新，操作结果以弹窗提示成功或失败'],
        ['例外处理', '药材名称为空时拒绝提交并提示；名称重复时拒绝添加并提示唯一性冲突'],
    ]
)

# 模块二
doc.add_heading('1.2.2 模块二：药材查询', level=3)
add_table(
    ['项目', '内容'],
    [
        ['用例名称', '多字段模糊查询与详情展示'],
        ['参与者', '系统用户'],
        ['前置条件', '数据库中已有药材数据'],
        ['主流程',
         '1. 用户输入关键词（支持中文或拼音）\n'
         '2. 在下拉框中选择搜索范围：名称 / 拼音 / 功效分类 / 功效 / 别名\n'
         '3. 系统执行 LIKE 模糊匹配并展示结果列表（ID/名称/拼音/分类/功效/性味/归经）\n'
         '4. 用户点击某条结果，下方详情面板展示该药材的全部10个属性'],
        ['后置条件', '搜索结果列表与详情面板同步更新'],
        ['扩展', '点击"显示全部"按钮可恢复浏览全部药材，搜索框支持回车键触发查询'],
    ]
)

# 模块三
doc.add_heading('1.2.3 模块三：配伍禁忌检查（核心业务）', level=3)
add_table(
    ['项目', '内容'],
    [
        ['用例名称', '多味药材配伍禁忌自动检测'],
        ['参与者', '系统用户'],
        ['前置条件', '数据库中有药材数据与禁忌规则数据'],
        ['主流程',
         '1. 用户从可选药材列表双击（或选中后点击添加）将药材加入待检查列表\n'
         '2. 可添加多味药材，支持移除和清空操作\n'
         '3. 点击"执行配伍禁忌检查"\n'
         '4. 系统对待检查列表中的药材进行两两组合（C(n,2)组合）\n'
         '5. 每组药材对名称经 sorted() 归一化后与禁忌规则库逐一比对\n'
         '6. 若存在冲突，输出每处冲突的药材对、禁忌类型（十八反/十九畏/其它）及详细说明\n'
         '   若无冲突，提示"未发现配伍禁忌"'],
        ['后置条件', '检查结果显示在结果文本区域'],
        ['例外处理', '待检查药材少于2味时弹出提示"请至少选择两味药材进行检查"'],
    ]
)

# 模块四
doc.add_heading('1.2.4 模块四：禁忌规则管理', level=3)
add_table(
    ['项目', '内容'],
    [
        ['用例名称', '配伍禁忌规则增删改查维护'],
        ['参与者', '系统用户'],
        ['前置条件', '数据库已初始化'],
        ['主流程',
         '1. 用户打开"禁忌规则管理"标签页，左侧展示全部规则列表\n'
         '2. 选中规则后右侧表单自动填充（药材A/B从下拉框加载全部药材名）\n'
         '3. 禁忌类型可选"十八反""十九畏""其它"\n'
         '4. 说明字段支持多行文本输入\n'
         '5. 新增/更新/删除操作均有弹窗确认与结果反馈'],
        ['约束规则',
         '药材A与药材B不能相同（前端校验）\n'
         '同一药材对的规则不能重复（数据库 UNIQUE 约束）\n'
         '存储时药材对自动按字母序排列以保证唯一性（Python 层 sorted()）'],
    ]
)

# 模块五
doc.add_heading('1.2.5 模块五：数据统计', level=3)
add_table(
    ['项目', '内容'],
    [
        ['用例名称', '系统数据概览统计'],
        ['参与者', '系统用户'],
        ['前置条件', '数据库已初始化并包含数据'],
        ['主流程',
         '1. 用户切换到"数据统计"标签页\n'
         '2. 点击"刷新统计"按钮\n'
         '3. 系统展示：药材总数、禁忌规则总数、各功效分类药材数量分布（GROUP BY category）、'
         '各禁忌类型规则数量分布（GROUP BY rule_type）'],
        ['数据来源', '通过 SQL 聚合查询 COUNT + GROUP BY 实时计算，保证数据时效性'],
    ]
)

# 模块六
doc.add_heading('1.2.6 模块六：数据操作', level=3)
add_table(
    ['项目', '内容'],
    [
        ['用例名称', 'JSON 导入导出与数据库重置'],
        ['参与者', '系统用户'],
        ['前置条件', '数据库已初始化'],
        ['主流程',
         '导出：用户选择目标路径 → 系统将全部药材与规则数据序列化为 JSON 文件'
         '（UTF-8 编码，indent=2 格式化）→ 弹窗提示导出条数\n'
         '导入：用户选择 JSON 文件 → 系统解析后逐条写入数据库（利用已有 add_herb/add_rule '
         '方法，自动跳过重复记录）→ 弹窗提示导入条数并刷新全部列表\n'
         '重置：二次确认弹窗 → DROP 全部表并重建 → 重新载入内置种子数据（25种药材+17条规则）'],
        ['数据格式', '{"herbs": [...], "incompatibilities": [...]}'],
    ]
)

# 1.3 非功能需求
doc.add_heading('1.3 非功能需求', level=2)
add_table(
    ['类别', '需求描述', '实现策略'],
    [
        ['技术约束', '零第三方依赖，仅使用 Python 标准库', 'sqlite3 + tkinter + json'],
        ['平台约束', '桌面 GUI 应用，跨平台可用', 'tkinter（Python 内置 GUI 工具包）'],
        ['数据持久化', '本地数据库存储，无需网络连接', 'SQLite3 文件数据库（tcm.db），嵌入式无服务'],
        ['数据完整性', '药材名称唯一、禁忌规则药材对唯一', 'UNIQUE 约束 + Python 层 sorted() 归一化'],
        ['可用性', '界面简洁直观，操作有即时反馈', '标签页分区设计，弹窗提示，列表实时刷新'],
        ['可测试性', '支持自动化测试，测试与生产数据隔离', 'tempfile 临时目录独立数据库，setup/teardown 隔离'],
        ['可维护性', '模块化分层，职责清晰', '数据层(database.py) / 表现层(ui.py) / 入口(main.py) 三层分离'],
        ['性能', '本地数据库操作响应时间 < 100ms', 'SQLite 文件数据库本地读写，配伍检查在 Python 内存中完成'],
    ]
)

# 1.4 用例图描述
doc.add_heading('1.4 用例总览', level=2)
add_para('系统包含单一参与者"系统用户"，与六大功能模块中的全部子用例存在关联关系。用例间的 include/extend 关系如下：', indent=True)
add_para('（1）"修改药材信息"和"删除药材" include "浏览药材列表"——均需先从列表中选中目标药材。')
add_para('（2）"查看药材详情" extend 各查询用例——用户从查询结果中点击某条记录触发详情的扩展行为。')
add_para('（3）"执行配伍禁忌检查" include "选择待检查药材"——检查前必须完成药材选择。')
add_para('（4）"修改/删除禁忌规则" include "浏览规则列表"——均需先从列表中选中目标规则。')
add_para('（5）"重置数据库" extend "导出数据到JSON"——重置前建议先导出备份。')
add_para('（完整的 PlantUML 用例图见 plantuml/scenarios_view.txt）', bold=False)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第二章：面向对象设计
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第二章  面向对象设计', level=1)

# 2.1 设计方法选择
doc.add_heading('2.1 设计方法选择', level=2)
add_para(
    '本项目采用面向对象设计方法（Object-Oriented Design），而非面向数据流的结构化方法（Structured Design）。'
    '选择理由如下：',
    indent=True
)
add_table(
    ['考量维度', '面向对象（本系统选择）', '面向数据流（未选择）'],
    [
        ['系统复杂度', '中等（2个核心类，~900行代码）', '更适合以变换为中心的系统'],
        ['数据与操作关系', '药材数据与其CRUD操作自然内聚在DBManager中', '数据与操作分离描述，破坏内聚性'],
        ['GUI交互', '事件驱动模型天然适合对象封装', '难以描述事件驱动的用户交互逻辑'],
        ['可扩展性', '新增功能只需扩展现有类或添加新类', '需修改数据流图全局结构'],
        ['Python语言特性', '原生支持类与对象，OOP是Python惯用范式', '函数式也可，但丧失封装优势'],
        ['测试友好性', '通过依赖注入可替换具体实现', '模块间耦合度高，难以独立测试'],
    ]
)

# 2.2 类的识别与职责分配
doc.add_heading('2.2 类的识别与职责分配', level=2)
add_para('从需求分析中提取名词与动词，采用"名词→类、动词→方法"的经典OOA策略，识别出以下核心类：')

doc.add_heading('2.2.1 DBManager（数据访问对象）', level=3)
add_para('数据访问层核心类，封装对 SQLite3 数据库的全部操作。', indent=True)
add_table(
    ['项目', '说明'],
    [
        ['所属模块', 'database.py'],
        ['职责',
         '封装对 SQLite3 数据库的全部 CRUD 操作、模糊查询、配伍禁忌检查、统计分析、'
         'JSON 导入导出、数据库重置与种子数据导入。为上层（表现层/测试层）提供语义化的方法接口，'
         '隐藏 SQL 实现细节。'],
        ['核心属性', 'conn: sqlite3.Connection — 数据库连接对象\n'
         'db_path: str — 数据库文件路径'],
        ['药材操作', 'add_herb(data: dict) → int\n'
         'update_herb(herb_id, data) → bool\n'
         'delete_herb(herb_id) → bool\n'
         'get_herb(herb_id) → dict | None\n'
         'get_all_herbs() → list[dict]\n'
         'get_herb_by_name(name) → dict | None'],
        ['查询操作', 'search_herbs(keyword: str, field: str) → list[dict]\n'
         '支持按 name / pinyin / category / gongxiao / alias 五个字段模糊查询'],
        ['规则操作', 'add_rule(herb_a, herb_b, rule_type, description) → int\n'
         'update_rule(rule_id, herb_a, herb_b, ...) → bool\n'
         'delete_rule(rule_id) → bool\n'
         'get_all_rules() → list[dict]'],
        ['核心业务', 'check_incompatibility(herb_names: list[str]) → list[dict]\n'
         '对输入药材名两两组合，与禁忌规则库逐一比对，返回所有冲突详情'],
        ['统计操作', 'get_statistics() → dict\n'
         '返回 herb_count、rule_count、categories 分布、rule_types 分布'],
        ['数据操作', 'export_to_json(filepath: str) → int\n'
         'import_from_json(filepath: str) → tuple[int, int]\n'
         'reset_database() — DROP + 重建表\n'
         'seed_data(herbs, rules) — 批量导入种子数据'],
        ['设计原则', '单一职责原则（SRP）：只负责数据访问\n'
         '信息隐藏：SQL语句封装在方法内部\n'
         '防御式设计：field白名单防注入、IntegrityError去重处理'],
    ]
)

doc.add_heading('2.2.2 TCMApp（GUI应用主类）', level=3)
add_para('表现层核心类，构建和管理全部GUI界面组件。', indent=True)
add_table(
    ['项目', '说明'],
    [
        ['所属模块', 'ui.py'],
        ['职责',
         '构建 6 个功能标签页的完整 GUI 界面（tkinter Notebook），处理用户交互事件，'
         '调用 DBManager 完成实际数据操作，将操作结果渲染回界面。'],
        ['核心属性',
         'db: DBManager — 数据层引用（通过构造函数注入）\n'
         'root: Tk — 主窗口对象\n'
         'herb_tree / search_tree / rule_tree: ttk.Treeview — 三个列表组件\n'
         'herb_fields: dict[str, tk.Text] — 药材编辑表单字段映射\n'
         'check_listbox / available_listbox: tk.Listbox — 配伍检查的双列表'],
        ['Tab构建方法',
         '_build_herb_manage_tab() — 药材管理（列表+表单+按钮）\n'
         '_build_herb_search_tab() — 药材查询（搜索栏+结果列表+详情面板）\n'
         '_build_incompatibility_check_tab() — 配伍检查（双列表+检查按钮+结果区）\n'
         '_build_rule_manage_tab() — 规则管理（列表+表单+按钮）\n'
         '_build_statistics_tab() — 数据统计（Text显示+刷新按钮）\n'
         '_build_data_ops_tab() — 数据操作（导入/导出/重置按钮）'],
        ['事件处理',
         '_add_herb / _update_herb / _delete_herb — 药材增删改\n'
         '_search_herbs / _on_search_select — 查询与详情展示\n'
         '_do_incompatibility_check / _add_to_check / _remove_from_check — 配伍检查\n'
         '_add_rule / _update_rule / _delete_rule — 规则管理\n'
         '_export_data / _import_data / _reset_database — 数据操作'],
        ['设计原则',
         '依赖倒置原则（DIP）：依赖 DBManager 抽象而非具体 SQL 实现\n'
         '开闭原则（OCP）：新增Tab只需添加一个 _build_xxx_tab() 方法\n'
         '每个Tab方法内聚封装该功能模块的全部UI组件与事件逻辑'],
    ]
)

# 2.3 类间关系
doc.add_heading('2.3 类之间的关系', level=2)
add_table(
    ['关系类型', '涉及类', '说明'],
    [
        ['关联（Association）', 'TCMApp → DBManager',
         'TCMApp 持有 DBManager 的引用，通过构造函数注入（依赖注入）'],
        ['依赖（Dependency）', 'main() → DBManager, TCMApp',
         'main() 函数作为装配器（Composition Root），负责创建对象并组装'],
        ['组合（Composition）', 'TCMApp → tkinter 子组件',
         'Treeview、Listbox、Text 等组件的生命周期与 TCMApp 绑定'],
    ]
)

add_para(
    '关键设计决策：TCMApp 不在自己的 __init__() 中创建 DBManager 实例，而是通过构造函数参数接收。'
    '这一"依赖注入"设计使得测试模块（test.py）可以向 TCMApp 传入指向临时数据库的 DBManager 实例，'
    '实现测试环境与生产数据的完全隔离。',
    indent=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第三章：软件体系结构风格与设计模式
# ═══════════════════════════════════════════════════════════════
doc.add_heading('第三章  软件体系结构风格与设计模式', level=1)

# 3.1 分层架构
doc.add_heading('3.1 体系结构风格：分层架构（Layered Architecture）', level=2)
add_para(
    '本系统采用严格的分层架构风格，将系统划分为三个层次。上层可以调用下层的公开接口，'
    '下层对上层完全无知。跨层通信必须经过中间层，不允许越层访问。',
    indent=True
)
add_table(
    ['层次', '对应模块', '职责', '依赖方向'],
    [
        ['表现层\n(Presentation)', 'ui.py — TCMApp',
         'GUI组件构建·事件绑定·表单校验·结果渲染', '依赖 ↓\n数据访问层'],
        ['数据访问层\n(Data Access)', 'database.py — DBManager',
         'CRUD封装·配伍逻辑·统计·导入导出·事务管理', '依赖 ↓\n数据层 + Python标准库'],
        ['数据层\n(Data)', 'tcm.db — SQLite3',
         'herbs表·incompatibilities表·数据持久化', '无依赖\n（被上层读写）'],
    ]
)

add_para('层次间交互规则的代码体现：', bold=True)

add_para('① main.py — 装配器（Composition Root），负责创建各层实例并完成依赖注入：')
add_code_block('''from database import DBManager   # 数据访问层
from ui import TCMApp             # 表现层

def main():
    db = DBManager()              # 创建数据访问层实例
    app = TCMApp(db)              # 注入到表现层（依赖注入）
    app.run()                     # 启动GUI事件循环
    db.close()''')

add_para('② ui.py 只依赖 DBManager 的公开方法，不直接操作 SQLite Connection：')
add_code_block('''class TCMApp:
    def __init__(self, db: DBManager):
        self.db = db   # 接收注入的数据层实例，内部不自行创建

    def _search_herbs(self):
        results = self.db.search_herbs(keyword, field)  # 调用数据层接口
        # 渲染结果到 Treeview...''')

add_para('③ database.py 对上层完全无知——DBManager 的代码中没有任何对 TCMApp 或 Tk 的引用。')

# 3.2 设计模式
doc.add_heading('3.2 设计模式应用', level=2)

# 3.2.1 DAO
doc.add_heading('3.2.1 DAO / Repository 模式', level=3)
add_para(
    '意图：将数据访问逻辑与业务/表现逻辑分离，提供统一的数据操作接口，隔离底层存储实现细节。',
    indent=True
)
add_para('在本系统中，DBManager 类是整个系统中唯一操作数据库的类。外部模块永远不直接编写 SQL 语句，所有数据访问都通过 DBManager 的语义化方法完成。')
add_para('代码示例 — DAO 实现（database.py）：', bold=True)
add_code_block('''class DBManager:
    def search_herbs(self, keyword: str, field: str = "name") -> list[dict]:
        """按指定字段模糊查询药材。调用者无需知道 LIKE 语法。"""
        allowed = {"name", "pinyin", "category", "gongxiao", "alias"}
        if field not in allowed:
            field = "name"         # 白名单校验，防止SQL注入
        pattern = f"%{keyword}%"
        rows = self.conn.execute(
            f"SELECT * FROM herbs WHERE {field} LIKE ? ORDER BY id",
            (pattern,)
        ).fetchall()
        return [dict(r) for r in rows]

    def check_incompatibility(self, herb_names: list[str]) -> list[dict]:
        """检查多味药材间的配伍禁忌。调用者无需知道规则的存储结构。"""
        results = []
        names = [n.strip() for n in herb_names if n.strip()]
        if len(names) < 2:
            return results
        rules = self.get_all_rules()
        for i in range(len(names)):
            for j in range(i + 1, len(names)):
                a, b = sorted([names[i], names[j]])
                for rule in rules:
                    ra, rb = sorted([rule["herb_a"], rule["herb_b"]])
                    if a == ra and b == rb:
                        results.append({
                            "herb_1": names[i],
                            "herb_2": names[j],
                            "rule_type": rule["rule_type"],
                            "description": rule["description"]
                        })
        return results''')

add_para('代码示例 — DAO 的消费者（ui.py），调用方无需编写任何 SQL：', bold=True)
add_code_block('''def _search_herbs(self):
    keyword = self.search_entry.get().strip()
    field = self.search_field.get()
    results = self.db.search_herbs(keyword, field)  # 一行调用完成查询
    for item in self.search_tree.get_children():
        self.search_tree.delete(item)
    for h in results:
        self.search_tree.insert("", "end", values=(
            h["id"], h["name"], h["pinyin"], h["category"],
            h["gongxiao"], h.get("xingwei", ""), h.get("guijing", "")
        ))''')

add_para(
    'DAO模式带来的好处：如果将来需要更换数据库（如从SQLite迁移到PostgreSQL），只需修改DBManager内部实现，'
    'ui.py和test.py无需任何改动。测试时可以轻松替换为指向临时文件的DBManager实例。',
    indent=True
)

# 3.2.2 依赖注入
doc.add_heading('3.2.2 依赖注入模式（Dependency Injection）', level=3)
add_para(
    '意图：由外部组装器（Composition Root）负责创建依赖对象并注入到使用者中，'
    '而非让使用者在内部自行创建依赖。消除类之间的硬耦合，提高可测试性。',
    indent=True
)

add_para('代码示例 — 正确做法：依赖由外部注入', bold=True)
add_code_block('''# main.py — 装配器（Composition Root）
db = DBManager()         # 创建依赖
app = TCMApp(db)         # 注入依赖

# ui.py — 接收依赖（而非主动创建）
class TCMApp:
    def __init__(self, db: DBManager):
        self.db = db        # 依赖由外部注入，不自行创建''')

add_para('代码示例 — 反模式对比（如果这样写就错了）：', bold=True)
add_code_block('''# 反模式：在类内部硬编码依赖创建
class TCMApp:
    def __init__(self):
        self.db = DBManager()  # 紧耦合！无法替换为测试数据库''')

add_para(
    '依赖注入在本系统中的关键价值：测试模块 test.py 可以向 TCMApp 传入指向临时目录数据库的'
    'DBManager 实例（见 test.py 第 25 行: db = DBManager(TEST_DB)），从而在不影响生产数据库'
    'tcm.db 的前提下完成全部自动化测试。',
    indent=True
)

# 3.2.3 模板方法变体
doc.add_heading('3.2.3 模板方法模式（Template Method）— 隐式变体', level=3)
add_para(
    '意图：定义一个操作的算法骨架，将具体步骤延迟到子步骤（方法）中实现。'
    '本系统未使用经典的继承式模板方法，而是采用了组合式的变体。',
    indent=True
)
add_para('代码示例 — _build_ui() 定义GUI构建骨架，各 _build_xxx_tab() 实现具体步骤：', bold=True)
add_code_block('''def _build_ui(self):
    """构建标签页界面 — 这是模板骨架。"""
    notebook = ttk.Notebook(self.root)
    notebook.pack(fill="both", expand=True, padx=5, pady=5)

    # 每个步骤委托给专门的方法 — 统一接口、独立实现
    self._build_herb_manage_tab(notebook)            # 步骤1：药材管理
    self._build_herb_search_tab(notebook)            # 步骤2：药材查询
    self._build_incompatibility_check_tab(notebook)  # 步骤3：配伍检查
    self._build_rule_manage_tab(notebook)            # 步骤4：规则管理
    self._build_statistics_tab(notebook)             # 步骤5：数据统计
    self._build_data_ops_tab(notebook)               # 步骤6：数据操作''')

add_para(
    '每个 _build_xxx_tab() 方法内部又遵循相同的微模板：创建 Frame → 构建组件（Treeview/Listbox/'
    'Button/Text）→ 绑定事件回调 → 注册到 notebook。这种一致性使代码易于理解和维护，新增一个Tab'
    '只需实现一个新方法并在 _build_ui() 中注册一行。',
    indent=True
)

# 3.3 总结
doc.add_heading('3.3 体系结构风格与设计模式总结', level=2)
add_table(
    ['风格/模式', '应用位置', '核心价值'],
    [
        ['分层架构', 'main / ui / database / tcm.db 三层分离', '关注点分离，降低耦合，便于独立测试与维护'],
        ['DAO / Repository', 'database.py — DBManager 类', '隐藏数据访问细节（SQL），提供统一接口，支持数据库迁移'],
        ['依赖注入', 'main() 创建 DBManager 并传入 TCMApp', '解除 ui 与 database 的硬耦合，支撑测试隔离策略'],
        ['模板方法（变体）', 'ui.py — _build_ui() 协调各 Tab 构建', '统一GUI构建流程，各Tab独立开发互不干扰'],
    ]
)

# ── 保存 ─────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '需求分析.docx')
doc.save(output_path)
print(f'文档已生成：{output_path}')
