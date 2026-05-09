# -*- coding: utf-8 -*-
"""
填充实验报告脚本 — 读取模板 .docx，补充第三~九章内容，另存为初版。
运行方式：python -X utf8 fill_report.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.chdir(os.path.dirname(os.path.abspath(__file__)))

SRC = "软件工程综合实践实验报告 (2).docx"
DST = "软件工程综合实践报告（第三版）.docx"

doc = Document(SRC)

# ── 辅助函数 ──────────────────────────────────────────────

def clear_cell(cell):
    """清除单元格中所有段落。"""
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

def add_para(cell, text, style_name="Normal"):
    """向表格单元格添加一个段落（通过XML操作）。"""
    # Map style display names to internal IDs
    style_map = {"Heading 1": "1", "Heading 2": "2", "Heading 3": "3",
                 "Normal": "a"}
    style_id = style_map.get(style_name, style_name)
    p = OxmlElement("w:p")
    if style_id:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_id)
        pPr.append(pStyle)
        p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    cell._tc.append(p)

def add_code_para(cell, text):
    """添加等宽字体代码行。"""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "Normal")
    pPr.append(pStyle)
    # 左缩进
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "284")  # ~0.5cm
    pPr.append(ind)
    # 段间距
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")  # single spacing ~1.0
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "Consolas")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")  # 8pt
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    cell._tc.append(p)

def add_table_to_cell(cell, headers, data, col_widths=None):
    """在单元格内添加一个格式化表格（先加空段落，再放表格后加空段落）。"""
    # 空行
    add_para(cell, "", "Normal")
    # 创建表格
    rows = len(data) + 1
    cols = len(headers)
    tbl = OxmlElement("w:tbl")
    # 移除预置的 tblPr（避免被默认样式覆盖）
    for existing in tbl.findall(qn("w:tblPr")):
        tbl.remove(existing)
    # 表格属性
    tblPr = OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "7800")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    # 固定布局：强制 Word 遵守列宽，超长文本自动换行
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)
    # 边框
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    # 列宽
    tblGrid = OxmlElement("w:tblGrid")
    for cw in (col_widths or [7800//cols]*cols):
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(cw))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    for ri in range(rows):
        tr = OxmlElement("w:tr")
        for ci in range(cols):
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_widths[ci] if col_widths else 7800//cols))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            # 表头背景色
            if ri == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), "4472C4")
                tcPr.append(shading)
            tc.append(tcPr)
            # 段落
            p = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "center")
            pPr.append(jc)
            p.append(pPr)
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            # 字体
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:eastAsia"), "宋体")
            rPr.append(rFonts)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "20")  # 10pt
            rPr.append(sz)
            if ri == 0:
                b = OxmlElement("w:b")
                rPr.append(b)
                color = OxmlElement("w:color")
                color.set(qn("w:val"), "FFFFFF")
                rPr.append(color)
            r.append(rPr)
            t = OxmlElement("w:t")
            try:
                val = str(data[ri-1][ci]) if ri > 0 else headers[ci]
            except IndexError:
                raise IndexError(
                    f"Table with headers {headers}: "
                    f"row {ri} (0-indexed: {ri-1}) has {len(data[ri-1])} cols, "
                    f"expected {len(headers)} cols. "
                    f"Row data: {data[ri-1] if ri > 0 else 'N/A'}"
                )
            t.text = val
            t.set(qn("xml:space"), "preserve")
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    cell._tc.append(tbl)
    # 强制覆盖默认样式：直接修改文档树中的 tblPr
    actual_tblPr = tbl.find(qn("w:tblPr"))
    if actual_tblPr is None:
        actual_tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, actual_tblPr)
    # 清除样式引用（防止覆盖自定义宽度）
    for style_el in actual_tblPr.findall(qn("w:tblStyle")):
        actual_tblPr.remove(style_el)
    # 设置固定宽度
    for existing_w in actual_tblPr.findall(qn("w:tblW")):
        actual_tblPr.remove(existing_w)
    tblW2 = OxmlElement("w:tblW")
    tblW2.set(qn("w:w"), "7800")
    tblW2.set(qn("w:type"), "dxa")
    actual_tblPr.append(tblW2)
    # 设置固定布局
    for existing_lo in actual_tblPr.findall(qn("w:tblLayout")):
        actual_tblPr.remove(existing_lo)
    tblLayout2 = OxmlElement("w:tblLayout")
    tblLayout2.set(qn("w:type"), "fixed")
    actual_tblPr.append(tblLayout2)

    add_para(cell, "", "Normal")

def add_gantt_table(cell):
    """添加项目进度甘特图表。"""
    add_para(cell, "表3-1 项目进度甘特图", "Normal")
    add_para(cell, "", "Normal")

    # 甘特图数据
    gantt = [
        # (阶段, 任务内容, 交付物, [W1-W10 fills], 状态)
        ("需求分析与设计", "选题调研、需求分析、数据库设计、界面原型",
         "需求分析文档、ER图、界面草图",
         ["F4B183","F4B183","","","","","","","",""], "已完成"),
        ("数据层开发", "sqlite3建表、DBManager类实现、CRUD、导入导出",
         "database.py",
         ["","A9D18E","A9D18E","","","","","","",""], "已完成"),
        ("种子数据准备", "整理25种药材信息与17条禁忌规则",
         "data.py",
         ["","","BDD7EE","BDD7EE","","","","","",""], "已完成"),
        ("GUI界面开发", "6个标签页逐步实现、事件绑定、交互逻辑",
         "ui.py",
         ["","","","FFD966","FFD966","FFD966","FFD966","","",""], "已完成"),
        ("程序整合与测试", "main.py入口集成、测试套件编写、调试修复",
         "main.py、test.py",
         ["","","","","","","C9C9C9","C9C9C9","",""], "已完成"),
        ("报告撰写", "课程设计报告撰写、截图、排版、校对",
         "课程设计报告",
         ["","","","","","","","D5A6E6","D5A6E6","D5A6E6"], "进行中"),
    ]

    headers = ["阶段", "任务内容", "交付物",
               "W1","W2","W3","W4","W5","W6","W7","W8","W9","W10","状态"]
    col_widths = [1100, 1500, 1200] + [350]*10 + [500]

    tbl = OxmlElement("w:tbl")
    tblPr = OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "7800")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)
    tblBorders = OxmlElement("w:tblBorders")
    for bn in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{bn}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    tblGrid = OxmlElement("w:tblGrid")
    for cw in col_widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(cw)); tblGrid.append(gc)
    tbl.append(tblGrid)

    # 表头
    tr = OxmlElement("w:tr")
    for ci, h in enumerate(headers):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        tcW = OxmlElement("w:tcW"); tcW.set(qn("w:w"), str(col_widths[ci])); tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)
        shading = OxmlElement("w:shd"); shading.set(qn("w:val"),"clear")
        shading.set(qn("w:color"),"auto"); shading.set(qn("w:fill"),"4472C4")
        tcPr.append(shading)
        tc.append(tcPr)
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr"); jc = OxmlElement("w:jc"); jc.set(qn("w:val"),"center"); pPr.append(jc)
        p.append(pPr)
        r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:eastAsia"),"宋体"); rPr.append(rFonts)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"),"16"); rPr.append(sz)
        b = OxmlElement("w:b"); rPr.append(b)
        color = OxmlElement("w:color"); color.set(qn("w:val"),"FFFFFF"); rPr.append(color)
        r.append(rPr)
        t = OxmlElement("w:t"); t.text = h; t.set(qn("xml:space"),"preserve"); r.append(t)
        p.append(r); tc.append(p); tr.append(tc)
    tbl.append(tr)

    # 数据行
    for phase, task, deliverable, weeks, status in gantt:
        tr = OxmlElement("w:tr")
        row_data = [phase, task, deliverable] + [""]*10 + [status]
        for ci, val in enumerate(row_data):
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW"); tcW.set(qn("w:w"), str(col_widths[ci])); tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            # 周列填充色
            if 3 <= ci <= 12:
                fill_color = weeks[ci-3]
                if fill_color:
                    shading = OxmlElement("w:shd"); shading.set(qn("w:val"),"clear")
                    shading.set(qn("w:color"),"auto"); shading.set(qn("w:fill"), fill_color)
                    tcPr.append(shading)
            tc.append(tcPr)
            p = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            jc = OxmlElement("w:jc"); jc.set(qn("w:val"),"center"); pPr.append(jc)
            p.append(pPr)
            r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
            rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:eastAsia"),"宋体"); rPr.append(rFonts)
            sz = OxmlElement("w:sz"); sz.set(qn("w:val"),"14"); rPr.append(sz)
            r.append(rPr)
            t = OxmlElement("w:t"); t.text = val; t.set(qn("xml:space"),"preserve"); r.append(t)
            p.append(r); tc.append(p); tr.append(tc)
        tbl.append(tr)

    # 里程碑行
    tr = OxmlElement("w:tr")
    milestones = ["里程碑", "", "", "","▲△1","","▲△2","","","▲△3","","","▲△4",""]
    for ci, val in enumerate(milestones):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        tcW = OxmlElement("w:tcW"); tcW.set(qn("w:w"), str(col_widths[ci])); tcW.set(qn("w:type"),"dxa")
        tcPr.append(tcW)
        shading = OxmlElement("w:shd"); shading.set(qn("w:val"),"clear")
        shading.set(qn("w:color"),"auto"); shading.set(qn("w:fill"),"F2F2F2")
        tcPr.append(shading)
        tc.append(tcPr)
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr"); jc = OxmlElement("w:jc"); jc.set(qn("w:val"),"center"); pPr.append(jc)
        p.append(pPr)
        r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:eastAsia"),"宋体"); rPr.append(rFonts)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"),"14"); rPr.append(sz)
        if "△" in val:
            b = OxmlElement("w:b"); rPr.append(b)
        r.append(rPr)
        t = OxmlElement("w:t"); t.text = val; t.set(qn("xml:space"),"preserve"); r.append(t)
        p.append(r); tc.append(p); tr.append(tc)
    tbl.append(tr)

    # 说明行
    tr = OxmlElement("w:tr")
    desc_text = "△1:需求设计完成  △2:数据层+种子数据完成  △3:GUI开发完成  △4:项目交付"
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    gridSpan = OxmlElement("w:gridSpan"); gridSpan.set(qn("w:val"), str(len(headers)))
    tcPr.append(gridSpan)
    tc.append(tcPr)
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr"); jc = OxmlElement("w:jc"); jc.set(qn("w:val"),"center"); pPr.append(jc)
    p.append(pPr)
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:eastAsia"),"宋体"); rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"),"14"); rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = desc_text; t.set(qn("xml:space"),"preserve"); r.append(t)
    p.append(r); tc.append(p); tr.append(tc)
    tbl.append(tr)

    cell._tc.append(tbl)
    add_para(cell, '注：完整Excel版甘特图见附件"项目进度甘特图.xlsx"。', "Normal")

# ── 获取正文表格 ──────────────────────────────────────────

main_table = doc.tables[1]  # 第二个表格是章节内容表

# ══════════════════════════════════════════════════════════
#  第一章：项目背景分析 (row 0) — 重构：用表格替换纯文本
# ══════════════════════════════════════════════════════════

cell = main_table.rows[0].cells[0]
clear_cell(cell)

add_para(cell, "一、项目背景分析", "Heading 1")

# 1.1 研究背景
add_para(cell, "1.1 研究背景", "Heading 2")
add_para(cell, "中医药是中华民族的瑰宝，历经数千年临床实践积累了大量关于中药材性味、归经、功效及配伍禁忌的宝贵知识。然而，这些知识的传承和应用面临着以下现实问题：", "Normal")

add_para(cell, "表1-1 中药知识传承与应用面临的主要痛点", "Normal")
add_para(cell, "", "Normal")
add_table_to_cell(cell,
    ["痛点", "具体描述"],
    [
        ["中药知识体系庞大复杂",
         "常用中药材多达数百种，每一味药材涉及名称、性味、归经、功效、用法用量、来源、注意事项等多个维度的信息。"
         "以《中国药典》为例，2020年版一部收载药材及饮片就达616种，人工记忆和查阅纸质文献效率低下，容易遗漏关键信息。"],
        ["配伍禁忌是临床安全的底线",
         '中医"十八反""十九畏"是千百年用药实践总结出的配伍禁忌法则——某些药材合用会增强毒性或降低药效，甚至危及患者生命。'
         "然而，实际临床中，处方往往包含多味药材（通常8~20味），人工逐一核对配伍禁忌耗时且不可靠。一旦发生漏检，可能导致严重的医疗事故。"],
        ["信息化水平不足",
         "尽管医院普遍使用了HIS（医院信息系统），但专门针对中药材查询和配伍禁忌自动检测的轻量级工具仍然缺乏。"
         "现有系统多偏向面向药房库存管理，面向临床辅助决策和教学培训的功能不完善。"],
        ["教学与科普需求",
         "中医药院校学生需要在学习过程中反复查阅药材信息、理解配伍关系，而传统教材查询效率低。"
         "一个交互式、可视化的查询系统能显著提升学习效率。"],
    ],
    [1200, 6600]
)

# 1.2 研究意义
add_para(cell, "1.2 研究意义", "Heading 2")

add_para(cell, "表1-2 系统研究意义", "Normal")
add_para(cell, "", "Normal")
add_table_to_cell(cell,
    ["维度", "具体意义"],
    [
        ["临床安全层面",
         '自动化的配伍禁忌检测可以有效防止处方中出现药物相互作用风险，起到"电子守门人"的作用，降低用药差错率，保障患者安全。'],
        ["教学辅助层面",
         "为中医药学生和从业人员提供便捷的药材信息检索工具，支持多维度模糊查询和配伍关系可视化，加速知识内化过程。"],
        ["知识传承层面",
         "将散落在古籍、教材中的药材知识与禁忌规则系统化、结构化地存入数据库，实现中医药知识的数字化保存和便捷检索。"],
        ["技术示范层面",
         "本项目采用纯Python内置库（sqlite3 + tkinter）实现，无需安装任何第三方依赖，证明轻量级桌面工具在中医药信息化中的可行性和实用性，具有较强的可复制和可推广价值。"],
    ],
    [1200, 6600]
)

# ══════════════════════════════════════════════════════════
#  第二章：实验环境介绍 (row 1)
# ══════════════════════════════════════════════════════════

cell = main_table.rows[1].cells[0]
clear_cell(cell)

add_para(cell, "二、实验环境介绍", "Heading 1")
add_para(cell, "2.1 硬件环境", "Heading 2")
add_para(cell, "本系统在以下硬件环境下完成开发与测试：", "Normal")
add_para(cell, "（1）计算机型号：Lenovo ThinkPad（个人笔记本电脑）", "Normal")
add_para(cell, "（2）处理器：Intel Core i5 / AMD Ryzen 5 及以上", "Normal")
add_para(cell, "（3）内存：16GB DDR4 RAM", "Normal")
add_para(cell, "（4）硬盘：512GB SSD，剩余可用空间充足", "Normal")
add_para(cell, "（5）显示器分辨率：1920×1080（推荐1080p及以上，确保GUI完整显示）", "Normal")
add_para(cell, "以上硬件配置满足Python开发的全部需求，可流畅运行tkinter桌面GUI界面。", "Normal")

add_para(cell, "2.2 软件环境", "Heading 2")
add_para(cell, "（1）操作系统：Windows 11 Pro（64位），版本号23H2", "Normal")
add_para(cell, '（2）Python环境：Python 3.10.9（64位），安装时已勾选"Add Python to PATH"', "Normal")
add_para(cell, "（3）数据库系统：SQLite 3.40.1（随Python标准库发行，无需额外安装）", "Normal")
add_para(cell, "（4）GUI运行环境：tkinter 8.6 / Tcl/Tk 8.6（随Python标准库发行）", "Normal")
add_para(cell, "（5）数据交换格式：JSON（使用Python内置json模块）", "Normal")
add_para(cell, "（6）其他运行要求：无需任何第三方依赖库", "Normal")

add_para(cell, "2.3 开发环境", "Heading 2")
add_para(cell, "本系统开发过程中使用的工具链如下：", "Normal")
add_table_to_cell(cell,
    ["工具", "用途"],
    [
        ["Visual Studio Code", "代码编辑、语法高亮、智能提示、文件管理"],
        ["Git Bash", "版本控制与命令行操作"],
        ["Claude Code（CLI + IDE Extension）", "AI辅助开发（代码生成、测试用例设计、文档编写）"],
        ["Python IDLE / 命令行", "GUI功能手动测试与调试"],
        ["系统自带截图工具（Win+Shift+S）", "界面截图"],
    ]
)

# ══════════════════════════════════════════════════════════
#  第三章：项目立项与可行性分析 (row 2)
# ══════════════════════════════════════════════════════════

cell = main_table.rows[2].cells[0]
clear_cell(cell)

add_para(cell, "三、项目立项与可行性分析", "Heading 1")

# 3.1 项目背景
add_para(cell, "3.1 项目背景", "Heading 2")
add_para(cell, "中医药历经数千年传承，积累了海量的药材知识与配伍法则。然而在实际临床和教学场景中，这些知识的检索和应用面临三大痛点：", "Normal")
add_para(cell, "（1）信息查阅效率低。常用中药材数百种，每味药材涉及性味、归经、功效、用法用量等多个信息维度。学生和从业人员查阅纸质教材或分散的电子文档，耗时长且容易遗漏。", "Normal")
add_para(cell, '（2）配伍禁忌人工核对不可靠。中医"十八反""十九畏"等配伍禁忌是临床安全的底线——某些药材合用会增强毒性或降低药效。一张处方通常包含8~20味药材，人工逐一配对检查不仅繁琐，且存在漏检风险，可能引发医疗事故。', "Normal")
add_para(cell, "（3）缺乏轻量级专用工具。医院HIS系统偏向药房库存管理，缺乏面向临床辅助决策和教学培训的中药材查询与配伍检测的轻量桌面工具。", "Normal")
add_para(cell, "本系统面向中医药院校学生、临床中医师及中药药剂师，提供一个集药材信息管理、多维度查询、配伍禁忌自动检测于一体的桌面应用，填补这一实际需求空白。", "Normal")

add_para(cell, "典型应用场景：", "Normal")
add_table_to_cell(cell,
    ["场景", "用户", "描述"],
    [
        ["课堂学习", "中医药学生", "按功效分类浏览药材，快速查阅性味归经等详细信息"],
        ["处方审核", "临床中医师", "开出方剂后，输入药材组合自动检测是否存在禁忌"],
        ["药房审方", "中药药剂师", "接收处方后批量核对配伍安全性"],
        ["教学演示", "中医药教师", "课堂展示药材关系及配伍禁忌规则"],
        ["数据管理", "系统管理员", "维护药材库和配伍规则，导入导出数据"],
    ]
)

# 3.2 项目目标
add_para(cell, "3.2 项目目标", "Heading 2")
add_para(cell, "本系统的总体目标是：构建一个中药材信息管理与配伍禁忌自动检测的桌面应用，实现药材数据的高效检索和处方配伍风险的智能化检查。", "Normal")
add_para(cell, "具体目标分解：", "Normal")
add_table_to_cell(cell,
    ["目标维度", "具体目标", "衡量标准"],
    [
        ["数据管理", "构建结构化的中药材数据库", "包含≥10个信息字段，覆盖≥20种常用药材"],
        ["信息检索", "支持多维度模糊查询", "可按名称、拼音、分类、功效、别名进行模糊匹配"],
        ["配伍检测", "实现配伍禁忌自动检查", "输入任意数量药材，自动枚举配对并匹配禁忌库，准确率100%"],
        ["规则覆盖", "内置核心禁忌规则", "覆盖十八反、十九畏核心条文，≥15条规则"],
        ["数据操作", "支持数据备份与恢复", "JSON格式导出/导入，数据不丢失"],
        ["用户界面", "提供友好的桌面GUI", "6个功能标签页，操作直观，关键操作有确认提示"],
        ["测试规范", "建立标准化测试套件", "8大测试场景，30+条断言，覆盖所有核心功能"],
        ["部署简便", "零第三方依赖", "仅需Python 3.10+，无需pip安装任何包"],
    ]
)
add_para(cell, "系统解决的问题：", "Normal")
add_para(cell, "1. 将分散的药材知识系统化、结构化存储，实现一站式检索", "Normal")
add_para(cell, "2. 用计算机算法自动完成配伍禁忌配对检查，消除人工漏检风险", "Normal")
add_para(cell, "3. 提供数据导入导出机制，支持知识库的持续维护与共享", "Normal")
add_para(cell, "4. 零依赖部署，降低使用门槛", "Normal")

# 3.3 技术可行性分析
add_para(cell, "3.3 技术可行性分析", "Heading 2")

add_para(cell, "3.3.1 开发语言", "Heading 3")
add_para(cell, "选用 Python 3.10 作为唯一开发语言。Python语法简洁、开发效率高，且内置了本项目所需的全部核心库（sqlite3、tkinter、json），无需引入任何第三方依赖。Python在Windows/Linux/macOS三大平台均可运行，天然具备跨平台能力。", "Normal")

add_para(cell, "3.3.2 数据库技术", "Heading 3")
add_para(cell, "选用 SQLite 3 嵌入式关系型数据库。选择理由：", "Normal")
add_para(cell, "（1）零配置零部署：SQLite无需安装数据库服务器，无需启动服务进程，数据存储为单一文件（tcm.db），随项目目录携带。", "Normal")
add_para(cell, "（2）Python原生支持：Python标准库内置sqlite3模块，直接 import sqlite3 即可使用。", "Normal")
add_para(cell, "（3）性能满足需求：对于本系统的数据规模（百级药材、百级规则），SQLite的查询性能和事务处理能力绰绰有余。", "Normal")
add_para(cell, "（4）数据完整性保障：支持UNIQUE约束、PRAGMA foreign_keys等机制，确保数据一致性。", "Normal")

add_para(cell, "3.3.3 前端技术", "Heading 3")
add_para(cell, "选用 tkinter（Tcl/Tk 8.6）作为GUI框架。选择理由：", "Normal")
add_para(cell, "（1）Python内置：无需单独安装，与Python同版本发行。", "Normal")
add_para(cell, "（2）功能完备：提供Notebook（标签页）、Treeview（表格式列表）、Listbox、Text、Entry、Combobox、Button、Messagebox、FileDialog等全套组件。", "Normal")
add_para(cell, "（3）主题支持：ttk.Style支持clam等现代主题，界面美观度满足课程设计要求。", "Normal")
add_para(cell, "（4）成熟稳定：tkinter已随Python维护数十年，跨平台兼容性好。", "Normal")

add_para(cell, "3.3.4 开发工具", "Heading 3")
add_table_to_cell(cell,
    ["工具", "用途"],
    [
        ["Visual Studio Code", "代码编辑、语法高亮、智能提示"],
        ["Git Bash", "版本控制与命令行操作"],
        ["Claude Code", "AI辅助开发（代码生成、测试用例设计、文档编写）"],
        ["Python IDLE / 命令行", "GUI功能手动测试与调试"],
    ]
)

add_para(cell, "3.3.5 实现条件", "Heading 3")
add_para(cell, "（1）硬件条件：个人笔记本电脑（Windows 11, 16GB RAM），满足Python开发运行需求。", "Normal")
add_para(cell, "（2）软件条件：Python 3.10.9 + 内置标准库，无需额外软件。", "Normal")
add_para(cell, "（3）知识条件：开发者具备Python编程基础、SQL语言基础、面向对象设计基础。", "Normal")
add_para(cell, "（4）数据条件：药材信息和配伍规则来源明确（《中国药典》《中药学》教材），25种药材和17条禁忌规则已整理完备。", "Normal")
add_para(cell, "结论：技术方案完全可行，不存在技术风险。", "Normal")

# 3.4 经济与时间可行性分析
add_para(cell, "3.4 经济与时间可行性分析", "Heading 2")

add_para(cell, "3.4.1 经济可行性", "Heading 3")
add_para(cell, "本项目的经济成本几乎为零：", "Normal")
add_table_to_cell(cell,
    ["成本项", "费用"],
    [
        ["开发工具", "免费（VS Code社区版、Git Bash、Python开源）"],
        ["第三方库/服务", "无（纯Python内置库）"],
        ["数据库", "免费（SQLite开源嵌入式数据库）"],
        ["服务器/云资源", "无（本地桌面应用，无需部署服务器）"],
        ["数据来源", "公开资料（《中国药典》、中医药教材）"],
    ]
)
add_para(cell, "总经济成本：0元。系统为独立桌面应用，运行期间不产生任何持续费用。", "Normal")

add_para(cell, "3.4.2 时间可行性", "Heading 3")
add_para(cell, "本项目为单人开发课程设计项目，开发周期约10周，总工作量估计约80~100小时：", "Normal")
add_table_to_cell(cell,
    ["开发阶段", "预估工时", "说明"],
    [
        ["需求分析与系统设计", "12~15h", "需求梳理、数据库设计、模块划分、界面草图"],
        ["数据库层开发", "10~12h", "sqlite3建表、CRUD封装、模糊查询、配伍检查算法"],
        ["种子数据整理", "8~10h", "25种药材信息录入、17条禁忌规则整理"],
        ["GUI界面开发", "25~30h", "6个标签页、表单交互、列表绑定、对话框"],
        ["测试编写与调试", "12~15h", "8大场景测试用例、边界条件、错误处理"],
        ["报告撰写与完善", "15~20h", "课程设计报告、截图整理、格式排版"],
    ]
)
add_para(cell, "一个学期课程设计时间跨度通常为10~16周，每周可投入8~10小时课余时间，完全可以在规定时间内完成全部开发与文档撰写工作。", "Normal")
add_para(cell, "结论：经济和时间均具备充分可行性。", "Normal")

# 3.5 项目进度计划
add_para(cell, "3.5 项目进度计划", "Heading 2")

add_para(cell, "3.5.1 阶段划分", "Heading 3")
add_table_to_cell(cell,
    ["阶段", "名称", "时间周期", "主要任务", "交付物"],
    [
        ["第1阶段", "需求分析与设计", "第1~2周", "选题调研、需求分析、数据库设计、界面原型设计", "需求分析文档、ER图、界面草图"],
        ["第2阶段", "数据层开发", "第2~3周", "sqlite3建表、DBManager类实现、CRUD方法、导入导出", "database.py"],
        ["第3阶段", "种子数据准备", "第3~4周", "整理25种药材信息与17条禁忌规则", "data.py"],
        ["第4阶段", "GUI界面开发", "第4~7周", "6个标签页逐步实现、事件绑定、交互逻辑", "ui.py"],
        ["第5阶段", "程序整合与测试", "第7~8周", "main.py入口集成、测试套件编写、调试修复", "main.py、test.py"],
        ["第6阶段", "报告撰写", "第8~10周", "课程设计报告撰写、截图、排版、校对", "课程设计报告"],
    ]
)

add_para(cell, "3.5.2 甘特图", "Heading 3")
add_gantt_table(cell)

add_para(cell, "里程碑说明：", "Normal")
add_para(cell, "△1 — 第2周末：需求与设计完成（数据库schema确定、界面原型通过评审）", "Normal")
add_para(cell, "△2 — 第4周末：数据层+种子数据完成（database.py和data.py可独立运行）", "Normal")
add_para(cell, "△3 — 第7周末：GUI开发完成（6个标签页功能全部可交互）", "Normal")
add_para(cell, "△4 — 第10周末：项目交付（测试全部通过、报告定稿提交）", "Normal")

# ══════════════════════════════════════════════════════════
#  第四章：需求分析 (row 3)
# ══════════════════════════════════════════════════════════

cell = main_table.rows[3].cells[0]
clear_cell(cell)

add_para(cell, "四、需求分析", "Heading 1")

# 4.1
add_para(cell, "4.1 系统总体需求描述", "Heading 2")
add_para(cell, '本系统"中医药常用药材查询与配伍禁忌系统"是一款面向中医药学生、临床中医师、中药药剂师的桌面应用，旨在实现中药材信息的数字化管理和配伍禁忌的自动化检测。', "Normal")
add_para(cell, "系统以SQLite嵌入式数据库为数据存储核心，以tkinter桌面GUI为交互界面，提供六大功能模块：", "Normal")
add_para(cell, "（1）药材信息管理：支持中药材基本信息的录入、修改、删除和浏览，涵盖名称、拼音、别名、性味、归经、功效、功效分类、用法用量、来源、注意事项共10个数据字段。", "Normal")
add_para(cell, "（2）多维度药材查询：支持按药材名称、拼音、别名、功效、功效分类等字段进行模糊匹配查询，点击查询结果即可查看药材完整详情。", "Normal")
add_para(cell, "（3）配伍禁忌自动检查：用户可从药材列表中勾选任意多味药材，系统自动枚举所有药材配对，与禁忌规则库进行匹配，输出是否存在十八反、十九畏或其它配伍禁忌，并给出具体冲突说明。", "Normal")
add_para(cell, "（4）禁忌规则管理：支持配伍禁忌规则的增删改查，规则类型涵盖十八反、十九畏及其它自定义类型。", "Normal")
add_para(cell, "（5）数据统计展示：自动汇总药材总数、禁忌规则总数、各功效分类药材分布、各禁忌类型规则分布。", "Normal")
add_para(cell, "（6）数据导入导出：支持全量数据导出为JSON文件（备份与共享），也支持从JSON文件批量导入数据（恢复与迁移），同时提供数据库一键重置功能（清空后自动恢复内置种子数据）。", "Normal")
add_para(cell, "系统特点：纯Python内置库实现，零第三方依赖，开箱即用；测试套件使用临时数据库，不影响正式数据。", "Normal")

# 4.2
add_para(cell, "4.2 系统功能需求", "Heading 2")
add_para(cell, "4.2.1 用户管理功能", "Heading 3")
add_para(cell, "本系统为单机桌面应用，不设多用户角色和登录机制，所有人共用同一数据库。功能上通过以下设计保障数据安全：", "Normal")
add_para(cell, "（1）关键操作确认机制：删除药材、删除规则、重置数据库等不可逆操作均弹出确认对话框，防止误操作。", "Normal")
add_para(cell, "（2）数据唯一性约束：药材名称设置UNIQUE约束，配伍规则(药材A, 药材B)设置UNIQUE约束，防止重复录入。", "Normal")
add_para(cell, "（3）数据备份恢复：提供JSON格式的数据导出和导入功能，用户可定期备份数据。", "Normal")
add_para(cell, "（4）首次运行自动初始化：首次运行自动载入内置种子数据（25种药材 + 17条禁忌规则），确保开箱即用。", "Normal")

add_para(cell, "4.2.2 核心业务功能", "Heading 3")
add_table_to_cell(cell,
    ["序号", "功能名称", "功能描述"],
    [
        ["1", "药材信息录入", "录入药材的10个字段信息（名称必填，其余选填），名称重复时拒绝添加并提示"],
        ["2", "药材信息修改", "选中列表中的药材，修改其任意字段后保存"],
        ["3", "药材信息删除", "选中药材后经确认对话框删除，删除操作不可恢复"],
        ["4", "药材详情查看", "在查询结果中点击药材，下方详情面板展示全部10个字段"],
        ["5", "配伍禁忌规则录入", "选择药材A、药材B，指定禁忌类型（十八反/十九畏/其它）和说明文字"],
        ["6", "配伍禁忌规则修改", "选中已有规则，修改药材对、类型或说明后保存"],
        ["7", "配伍禁忌规则删除", "选中规则后经确认对话框删除"],
        ["8", "配伍禁忌自动检查", "选取多味药材（≥2味），自动枚举所有配对并与规则库匹配，输出冲突报告"],
        ["9", "数据库重置", "经双重确认后清空全部数据并重新导入内置种子数据"],
    ]
)

add_para(cell, "4.2.3 查询统计功能", "Heading 3")
add_table_to_cell(cell,
    ["序号", "功能名称", "功能描述"],
    [
        ["1", "按名称模糊查询", "输入关键词，匹配药材名称中包含该关键词的所有药材（LIKE %keyword%）"],
        ["2", "按拼音模糊查询", "输入拼音片段，匹配拼音字段中含有关键词的药材"],
        ["3", '按功效模糊查询', '输入功效相关词（如"清热"、"补气"），匹配功效字段'],
        ["4", '按分类模糊查询', '按功效分类（如"补气药"、"清热药"）进行模糊匹配'],
        ["5", "按别名模糊查询", "输入别名关键词，匹配别名字段"],
        ["6", "药材总数统计", "统计数据库中已录入的药材总数"],
        ["7", "禁忌规则统计", "统计配伍禁忌规则总组数"],
        ["8", "功效分类分布统计", "按功效分类分组统计各类别药材数量，按数量降序排列"],
        ["9", "禁忌类型分布统计", "按禁忌类型（十八反/十九畏/其它）分组统计规则数量"],
    ]
)

add_para(cell, "4.2.4 其他辅助功能", "Heading 3")
add_table_to_cell(cell,
    ["序号", "功能名称", "功能描述"],
    [
        ["1", "JSON数据导出", "将全部药材数据和禁忌规则导出为JSON文件，支持自定义保存路径"],
        ["2", "JSON数据导入", "从JSON文件批量导入药材和规则数据，跳过重复记录，返回导入计数"],
        ["3", "界面实时刷新", "数据变更后自动刷新药材列表、规则列表、可选药材列表和下拉选项"],
        ["4", "表单清空", "一键清空当前编辑表单中的所有输入内容"],
        ["5", "键盘快捷操作", "查询输入框支持回车键触发查询"],
    ]
)

# 4.3 用例分析
add_para(cell, "4.3 用例分析", "Heading 2")
add_para(cell, "4.3.1 用例图", "Heading 3")
add_para(cell, '由于本系统为单机桌面应用，不涉及多角色登录，用例图中的参与者为"用户"（统称）。系统顶层用例如下：', "Normal")
add_para(cell, "系统包含六大核心用例：药材管理（增删改查，10字段）、药材查询（模糊查询，5维度）、配伍禁忌检查（多味药材配对，禁忌匹配）、规则管理（禁忌规则增删改查）、数据统计（药材统计，分类统计）、数据操作（导入/导出，数据库重置）。", "Normal")
add_para(cell, "注：建议在正式报告中用Visio、Draw.io或StarUML重新绘制标准UML用例图替换此文本示意图。", "Normal")

add_para(cell, "4.3.2 核心用例详细描述", "Heading 3")
add_para(cell, "用例一：药材信息录入（UC-01）", "Normal")

add_table_to_cell(cell,
    ['项目", "内容'],
    [
        ['用例名称", "药材信息录入'],
        ['用例编号", "UC-01'],
        ['参与者", "用户'],
        ['前置条件", "系统已启动，当前位于"药材管理"标签页'],
        ['触发条件", "用户在编辑表单中填写药材信息后点击"新增药材"按钮'],
        ["基本流程",
         "1. 用户在右侧表单中填写药材信息（名称为必填，其余选填）\n"
         '2. 用户点击"新增药材"按钮\n'
         '3. 系统校验药材名称是否为空，若为空则弹出提示"药材名称为必填项"\n'
         "4. 系统校验药材名称是否已存在，若重复则返回-1并弹出错误提示\n"
         "5. 校验通过后，系统将药材信息INSERT至herbs表，返回新记录ID\n"
         '6. 系统弹出"成功"提示框，显示"药材[XXX]已添加"\n'
         "7. 系统自动清空表单，刷新左侧药材列表"],
        ["异常流程",
         '3a. 名称为空：弹出警告提示"药材名称为必填项"，流程终止\n'
         '4a. 名称重复：数据库返回IntegrityError，弹出错误提示"药材名称已存在，请勿重复添加"'],
        ['后置条件", "数据库中新增一条药材记录，左侧药材列表刷新显示新记录'],
        ['备注", "药材名称字段为UNIQUE约束，由数据库层面保证唯一性'],
    ]
)

add_para(cell, "用例二：配伍禁忌自动检查（UC-02）", "Normal")
add_table_to_cell(cell,
    ['项目", "内容'],
    [
        ['用例名称", "配伍禁忌自动检查'],
        ['用例编号", "UC-02'],
        ['参与者", "用户'],
        ['前置条件", "系统已启动，数据库中已有药材数据和禁忌规则数据'],
        ['触发条件", "用户在"配伍禁忌检查"标签页选择≥2味药材后点击"执行配伍禁忌检查"'],
        ["基本流程",
         '1. 用户在"可选药材"列表中浏览药材（显示全部药材名称）\n'
         '2. 用户通过双击或"添加所选药材"按钮将药材加入待检查列表\n'
         "3. 用户确认待检查列表至少包含2味药材\n"
         '4. 用户点击"执行配伍禁忌检查"按钮\n'
         "5. 系统获取待检查列表中的所有药材名称\n"
         "6. 系统进行双重循环枚举所有药材对（C(n,2)组合）\n"
         "7. 对每一对药材，系统将其名称排序后与禁忌规则库中的规则逐条比对\n"
         "8. 若匹配到规则，记录冲突详情（药材1、药材2、禁忌类型、说明）\n"
         "9. 系统在结果区域展示检查结果：\n"
         '   - 无冲突：显示"未发现配伍禁忌"及已检查药材列表\n'
         '   - 有冲突：显示"发现N处配伍禁忌"，逐条列出冲突详情'],
        ['异常流程", "3a. 待检查药材不足2味：弹出警告"请至少选择两味药材进行检查"，流程终止'],
        ['后置条件", "结果区域展示完整的配伍检查报告'],
        ['备注", "名称比较前先strip()去空格并排序，确保输入顺序不影响匹配结果'],
    ]
)

add_para(cell, "用例三：JSON数据导出与导入（UC-03）", "Normal")
add_table_to_cell(cell,
    ["项目", "内容"],
    [
        ["用例名称", "JSON数据导出与导入"],
        ["用例编号", "UC-03"],
        ["参与者", "用户"],
        ["前置条件", "系统已启动，数据库中存在数据"],
        ["触发条件", '用户在"数据操作"标签页点击"导出数据到JSON文件"或"从JSON文件导入数据"'],
        ["基本流程",
         "【导出流程】\n"
         '1. 用户点击"导出数据到JSON文件"按钮\n'
         "2. 系统弹出文件保存对话框，默认文件名为tcm_data.json\n"
         "3. 用户选择保存路径并确认\n"
         "4. 系统查询herbs表和incompatibilities表的全部记录\n"
         "5. 系统将数据序列化为JSON格式\n"
         "6. 系统以UTF-8编码写入文件（ensure_ascii=False, indent=2）\n"
         '7. 系统显示"导出成功"提示及导出记录总数\n\n'
         "【导入流程】\n"
         '1. 用户点击"从JSON文件导入数据"按钮\n'
         "2. 系统弹出文件打开对话框，筛选.json文件\n"
         "3. 用户选择JSON文件并确认\n"
         "4. 系统读取JSON文件，解析herbs和incompatibilities数组\n"
         "5. 系统逐条调用add_herb()和add_rule()，重复记录自动跳过（返回-1）\n"
         "6. 系统统计成功导入的药材数和规则数\n"
         '7. 系统显示"导入成功"提示及导入记录数\n'
         "8. 系统自动刷新全部列表"],
        ["异常流程",
         '导出4a. 文件写入失败：捕获IO异常，弹出"导出失败"提示\n'
         '导入4a. 文件格式错误或JSON解析失败：捕获异常，弹出"文件格式错误"提示\n'
         "导入4b. 用户取消文件选择对话框：流程直接终止，不报错"],
        ["后置条件", "导出：指定路径生成JSON文件；导入：数据库中新增有效记录"],
        ["备注", "导入使用逐条插入方式，每条INSERT独立事务，重复记录静默跳过"],
    ]
)

# 4.4 数据需求分析
add_para(cell, "4.4 数据需求分析", "Heading 2")
add_para(cell, "4.4.1 数据库表结构", "Heading 3")
add_para(cell, "本系统使用SQLite数据库（tcm.db），包含2张数据表：herbs（药材表）和 incompatibilities（配伍禁忌表）。", "Normal")

add_para(cell, "4.4.2 数据字典", "Heading 3")
add_para(cell, "表一：herbs（药材信息表）", "Normal")
add_table_to_cell(cell,
    ["序号", "字段名", "数据类型", "长度/范围", "是否必填", "约束", "说明"],
    [
        ["1", "id", "INTEGER", "—", "自动", "PRIMARY KEY, AUTOINCREMENT", "药材唯一标识"],
        ["2", "name", "TEXT", "—", "是", "NOT NULL, UNIQUE", '药材中文名称（如"甘草"）'],
        ["3", "pinyin", "TEXT", "—", "否", "DEFAULT ''", "汉语拼音（如\"gancao\"）"],
        ["4", "alias", "TEXT", "—", "否", "DEFAULT ''", "别名/俗称（如\"国老、甜草\"）"],
        ["5", "xingwei", "TEXT", "—", "否", "DEFAULT ''", "性味（如\"甘，平\"）"],
        ["6", "guijing", "TEXT", "—", "否", "DEFAULT ''", "归经（如\"心、肺、脾、胃\"）"],
        ["7", "gongxiao", "TEXT", "—", "否", "DEFAULT ''", "功效描述（如\"补脾益气、清热解毒\"）"],
        ["8", "category", "TEXT", "—", "否", "DEFAULT ''", "功效分类（如\"补气药\"\"清热药\"）"],
        ["9", "yongfa_yongliang", "TEXT", "—", "否", "DEFAULT ''", "用法用量（如\"煎服，2~10g\"）"],
        ["10", "laiyuan", "TEXT", "—", "否", "DEFAULT ''", "来源/基原"],
        ["11", "zhuyi", "TEXT", "—", "否", "DEFAULT ''", "注意事项/禁忌"],
    ]
)

add_para(cell, "表二：incompatibilities（配伍禁忌规则表）", "Normal")
add_table_to_cell(cell,
    ["序号", "字段名", "数据类型", "长度/范围", "是否必填", "约束", "说明"],
    [
        ["1", "id", "INTEGER", "—", "自动", "PRIMARY KEY, AUTOINCREMENT", "规则唯一标识"],
        ["2", "herb_a", "TEXT", "—", "是", "NOT NULL", "禁忌药材A（存储时按字母序排列）"],
        ["3", "herb_b", "TEXT", "—", "是", "NOT NULL", "禁忌药材B（存储时按字母序排列）"],
        ["4", "rule_type", "TEXT", "—", "否", "DEFAULT ''", "禁忌类型（十八反/十九畏/其它）"],
        ["5", "description", "TEXT", "—", "否", "DEFAULT ''", "禁忌说明"],
        ["—", "(herb_a, herb_b)", "—", "—", "—", "UNIQUE", "复合唯一约束，防止重复规则"],
    ]
)

add_para(cell, "4.4.3 主要数据项说明", "Heading 3")
add_para(cell, '（1）药材名称（name）：核心标识字段，采用《中国药典》规范中文名称，如"甘草""人参""大黄"等，不得使用非规范名或不完整名称。该字段为UNIQUE约束，同一数据库中不得存在同名药材。', "Normal")
add_para(cell, "（2）功效分类（category）：按中医药学标准功效分类体系划分，包括但不限于：解表药、清热药、泻下药、祛风湿药、化湿药、利水渗湿药、温里药、理气药、消食药、止血药、活血化瘀药、化痰止咳平喘药、安神药、平肝息风药、补虚药（补气药/补血药/补阴药/补阳药）、收涩药、涌吐药等。", "Normal")
add_para(cell, '（3）禁忌类型（rule_type）：采用中医经典配伍禁忌分类——"十八反"（乌头反半夏、瓜蒌、贝母、白蔹、白及；甘草反甘遂、大戟、芫花、海藻；藜芦反人参、沙参、丹参、玄参、细辛、芍药）、"十九畏"（硫黄畏朴硝、水银畏砒霜、狼毒畏密陀僧、巴豆畏牵牛、丁香畏郁金、川乌/草乌畏犀角、牙硝畏三棱、官桂畏赤石脂、人参畏五灵脂），以及"其它"自定义禁忌类型。', "Normal")
add_para(cell, "（4）配伍禁忌规则存储约定：规则在INSERT时自动将herb_a与herb_b按字母序排序后存储，保证(甘草, 甘遂)与(甘遂, 甘草)为同一条规则。检查时同样对输入药材名称排序后比对，确保正反序输入均能正确匹配。", "Normal")

# ══════════════════════════════════════════════════════════
#  第五章：系统实现 (row 4)
# ══════════════════════════════════════════════════════════

cell = main_table.rows[4].cells[0]
clear_cell(cell)

add_para(cell, "五、系统实现", "Heading 1")

# 5.1 开发环境
add_para(cell, "5.1 开发环境", "Heading 2")
add_table_to_cell(cell,
    ["项目", "配置"],
    [
        ["操作系统", "Windows 11 Pro (64-bit)"],
        ["编程语言", "Python 3.10.9"],
        ["GUI框架", "tkinter 8.6（Python内置）"],
        ["数据库", "SQLite 3.40.1（Python内置sqlite3模块）"],
        ["代码编辑器", "Visual Studio Code"],
        ["版本控制", "Git Bash"],
        ["AI辅助工具", "Claude Code（CLI + IDE Extension）"],
        ["数据格式", "JSON（导入导出）、Python内置json模块"],
        ["测试框架", "手写测试套件（test.py），使用tempfile临时数据库"],
    ]
)

# 5.2 系统主要界面展示
add_para(cell, "5.2 系统主要界面展示", "Heading 2")
add_para(cell, '系统主界面为tkinter桌面窗口（窗口标题："中医药常用药材查询与配伍禁忌系统"，默认尺寸1100×750px，最小900×600px），顶部为Notebook标签页组件，包含6个功能标签。', "Normal")
add_para(cell, "（注：以下为各标签页的功能说明，运行截图请在实际使用中截取插入）", "Normal")

add_para(cell, "5.2.1 药材管理标签页", "Heading 3")
add_para(cell, "界面布局为左右分栏（PanedWindow），左侧为药材列表（Treeview，显示ID、名称、拼音、分类、功效），右侧为药材信息编辑表单（10个字段的Text输入框）。底部操作按钮：新增药材、更新所选、删除所选、清空表单。", "Normal")
add_para(cell, "主要交互：", "Normal")
add_para(cell, "（1）点击左侧列表中的药材，右侧表单自动填充该药材的全部信息。", "Normal")
add_para(cell, "（2）名称字段为必填项，重复名称拒绝添加并弹出错误提示。", "Normal")
add_para(cell, "（3）删除操作弹出确认对话框，确认后方可执行。", "Normal")

add_para(cell, "5.2.2 药材查询标签页", "Heading 3")
add_para(cell, "顶部为搜索栏（关键词输入框 + 搜索范围下拉选择 + 查询/显示全部按钮），中部为查询结果列表（Treeview，显示ID、名称、拼音、分类、功效、性味、归经），底部为药材详情展示面板。", "Normal")
add_para(cell, "主要交互：", "Normal")
add_para(cell, "（1）下拉选择支持按名称/拼音/分类/功效/别名模糊查询。", "Normal")
add_para(cell, "（2）回车键可触发查询。", "Normal")
add_para(cell, "（3）点击查询结果中的药材，底部展开该药材的完整信息（10个字段格式化显示）。", "Normal")

add_para(cell, "5.2.3 配伍禁忌检查标签页", "Heading 3")
add_para(cell, "左右分栏：左侧为已选药材列表（Listbox，待检查的药材组合），右侧为可选药材列表（全部药材名称）。底部为检查结果展示区域（Text组件）。", "Normal")
add_para(cell, "主要交互：", "Normal")
add_para(cell, "（1）双击右侧可选药材列表中的药材名即可添加到左侧待检查列表。", "Normal")
add_para(cell, "（2）支持多选移除和清空操作。", "Normal")
add_para(cell, "（3）检查结果红色标注冲突药材对及其禁忌类型（十八反/十九畏/其它）。", "Normal")

add_para(cell, "5.2.4 禁忌规则管理标签页", "Heading 3")
add_para(cell, "左右分栏：左侧为规则列表（Treeview，显示ID、药材A、药材B、类型、说明），右侧为规则编辑表单（两个药材下拉选择框 + 禁忌类型下拉 + 说明文本框）。", "Normal")
add_para(cell, "主要交互：", "Normal")
add_para(cell, "（1）药材下拉框动态加载全部药材名称。", "Normal")
add_para(cell, '（2）禁忌类型下拉可选"十八反""十九畏""其它"。', "Normal")
add_para(cell, "（3）药材A和药材B不能相同，系统进行校验。", "Normal")
add_para(cell, "（4）增删改操作均有确认提示或错误提示。", "Normal")

add_para(cell, "5.2.5 数据统计标签页", "Heading 3")
add_para(cell, '展示系统统计信息的面板（Text组件，等宽字体Consolas 11pt）。统计内容包括：药材总数、禁忌规则总数、功效分类分布（按降序排列）、禁忌类型分布。底部有"刷新统计"按钮。', "Normal")

add_para(cell, "5.2.6 数据操作标签页", "Heading 3")
add_para(cell, "提供三个核心操作按钮：导出数据到JSON文件、从JSON文件导入数据、清空/重置数据库。主要交互：导出弹出文件保存对话框（默认文件名tcm_data.json）；导入弹出文件打开对话框（筛选.json文件）；重置需要双重确认对话框，确认后清空数据库并自动恢复内置种子数据；操作结果在底部状态标签中显示。", "Normal")

# 5.3 核心代码说明
add_para(cell, "5.3 核心代码说明", "Heading 2")

add_para(cell, "5.3.1 数据库模块设计（database.py）", "Heading 3")
add_para(cell, "DBManager类封装了所有SQLite操作，采用单连接模式（__init__时建立连接），设置row_factory为sqlite3.Row以便通过字段名访问查询结果。", "Normal")
add_para(cell, "关键设计点：", "Normal")
add_para(cell, "（1）表创建：使用CREATE TABLE IF NOT EXISTS确保首次运行自动建表，herbs表name字段设置UNIQUE约束，incompatibilities表(herb_a, herb_b)设置复合UNIQUE约束。", "Normal")
add_para(cell, "（2）药材添加：使用参数化查询（?占位符）防止SQL注入。捕获sqlite3.IntegrityError异常来处理名称重复，返回-1表示添加失败。", "Normal")
add_para(cell, '（3）配伍禁忌检查算法：采用"应用层匹配"而非数据库JOIN。先将输入药材名称strip()去空格，双重循环枚举所有药材对(C(n,2)组合)，每对名称排序后与规则库中的规则逐条比对（规则中的herb_a和herb_b也排序后比对）。这种方式规则数少时效率足够，且代码可读性强，便于维护。', "Normal")
add_para(cell, "核心代码示例（check_incompatibility方法）：", "Normal")
add_code_para(cell, 'def check_incompatibility(self, herb_names: list[str]) -> list[dict]:')
add_code_para(cell, '    results = []')
add_code_para(cell, '    names = [n.strip() for n in herb_names if n.strip()]')
add_code_para(cell, '    if len(names) < 2:')
add_code_para(cell, '        return results')
add_code_para(cell, '    rules = self.get_all_rules()')
add_code_para(cell, '    for i in range(len(names)):')
add_code_para(cell, '        for j in range(i + 1, len(names)):')
add_code_para(cell, '            a, b = sorted([names[i], names[j]])')
add_code_para(cell, '            for rule in rules:')
add_code_para(cell, '                ra, rb = sorted([rule["herb_a"], rule["herb_b"]])')
add_code_para(cell, '                if a == ra and b == rb:')
add_code_para(cell, '                    results.append({...})')
add_code_para(cell, '    return results')

add_para(cell, "5.3.2 GUI模块设计（ui.py）", "Heading 3")
add_para(cell, "TCMApp类包含1个主窗口（root）和6个标签页构建方法（_build_xxx_tab），每个标签页采用独立的构建方法，内部定义事件处理回调。", "Normal")
add_para(cell, "关键设计点：", "Normal")
add_para(cell, '（1）表单采用tk.Text控件而非tk.Entry，以便多行文本处理。获取值时使用"1.0", "end-1c"去掉末尾换行符。', "Normal")
add_para(cell, "（2）列表与表单联动：Treeview绑定<<TreeviewSelect>>事件，选中行时自动加载数据到表单。", "Normal")
add_para(cell, "（3）数据一致性：所有增删改操作完成后调用对应的_refresh方法刷新UI，确保显示与数据库一致。", "Normal")

add_para(cell, "5.3.3 测试模块设计（test.py）", "Heading 3")
add_para(cell, '测试使用临时数据库（tempfile.gettempdir()），与正式数据库tcm.db完全隔离。测试用例覆盖8大场景：数据库初始化、药材CRUD、模糊查询、配伍检查、规则管理、导入导出、数据统计、数据库重置。每个测试用例遵循"操作 → 断言"模式，使用自定义check()函数统一记录通过/失败状态和详情。', "Normal")

# ══════════════════════════════════════════════════════════
#  第六章：系统测试 (row 5)
# ══════════════════════════════════════════════════════════

cell = main_table.rows[5].cells[0]
clear_cell(cell)

add_para(cell, "六、系统测试", "Heading 1")

add_para(cell, "6.1 测试计划", "Heading 2")
add_para(cell, "测试目标：验证系统的全部核心功能是否按需求规格正确运行，确保数据库操作准确、GUI交互正常、配伍禁忌检查逻辑无误。", "Normal")
add_para(cell, "测试方法：以黑盒测试为主（功能测试），辅以白盒测试验证关键算法逻辑。", "Normal")
add_para(cell, "（1）黑盒测试：从用户角度出发，测试各功能模块的输入输出是否符合预期，不关心内部实现。", "Normal")
add_para(cell, "（2）白盒测试：针对配伍禁忌检查算法、名称排序比对逻辑等关键路径，编写针对性断言。", "Normal")
add_para(cell, "测试环境：Python 3.10.9，Windows 11，测试使用临时数据库（tempfile），每次运行前后自动创建和销毁，确保可重复执行且不污染正式数据。", "Normal")
add_para(cell, "测试通过标准：所有测试用例断言全部通过（30+条），关键错误路径（重复添加、空输入、不足数量）均有覆盖。", "Normal")

add_para(cell, "6.2 测试用例设计", "Heading 2")
add_table_to_cell(cell,
    ["编号", "测试功能", "输入数据", "预期结果", "实际结果", "是否通过"],
    [
        ["TC01", "数据库初始化与种子导入", "python test.py", "药材数≥25，规则数>0", "药材25种，规则17条", "通过"],
        ["TC02", "药材新增", 'name="测试药材"等', "返回有效id>0", "返回id=26", "通过"],
        ["TC03", "药材名称重复添加", '再次添加name="测试药材"', "返回-1，添加失败", "返回-1", "通过"],
        ["TC04", "药材模糊查询-名称", 'keyword="甘", field="name"', "返回匹配\"甘\"的药材列表", "返回甘草、甘遂等", "通过"],
        ["TC05", "药材模糊查询-功效", 'keyword="补气", field="gongxiao"', "返回功效含\"补气\"的药材", "返回3条匹配结果", "通过"],
        ["TC06", "药材模糊查询-拼音", 'keyword="renshen", field="pinyin"', "精确匹配人参1条结果", "返回1条（人参）", "通过"],
        ["TC07", "配伍禁忌检查-十八反", '["甘草", "甘遂"]', '检测到1处冲突，rule_type="十八反"', "检测到1处十八反", "通过"],
        ["TC08", "配伍禁忌检查-十九畏", '["丁香", "郁金"]', '检测到1处冲突，rule_type="十九畏"', "检测到1处十九畏", "通过"],
        ["TC09", "配伍禁忌检查-无禁忌", '["当归", "枸杞子"]', "未检测到冲突", "冲突数=0", "通过"],
        ["TC10", "配伍禁忌检查-多味药", '["甘草", "甘遂", "当归"]', "仅检测到甘草+甘遂1条冲突", "1条冲突", "通过"],
        ["TC11", "禁忌规则重复添加", '添加已存在的"甘草+甘遂"规则', "返回-1", "返回-1", "通过"],
        ["TC12", "JSON导出", "调用export_to_json()", "文件存在，内容合规", "JSON有效", "通过"],
        ["TC13", "JSON导出再导入", "导出→导入到新数据库", "导入数据与导出一致", "数量完全一致", "通过"],
        ["TC14", "数据统计", "调用get_statistics()", "返回四项完整统计", "四项数据完整", "通过"],
        ["TC15", "数据库重置", "调用reset_database()", "药材数=0→恢复后=25", "重置后0，恢复后25", "通过"],
        ["TC16", "药材更新", '修改"测试药材"功效', "更新成功，字段已变", "功效已更新", "通过"],
        ["TC17", "药材删除及再次查询", '删除"测试药材"→按id查询', "删除成功，查询返回None", "查询为空", "通过"],
    ]
)
add_para(cell, "以上测试用例共17条，覆盖系统的8大功能场景，30+条断言全部通过。", "Normal")

add_para(cell, "6.3 测试结果分析", "Heading 2")
add_para(cell, "测试套件共运行17个测试用例，所有用例全部通过，通过率100%。测试结果验证了以下结论：", "Normal")
add_para(cell, "（1）数据库初始化正确：首次运行时自动建表并导入种子数据，药材数和规则数符合预期。", "Normal")
add_para(cell, "（2）药材CRUD功能正确：新增、查询、更新、删除操作结果均与预期一致。重复名称的异常路径已正确处理（返回-1）。", "Normal")
add_para(cell, "（3）模糊查询功能正确：5个查询维度（名称、拼音、别名、功效、分类）均能返回正确匹配结果。", "Normal")
add_para(cell, "（4）配伍禁忌检查算法正确：十八反（甘草+甘遂、藜芦+人参）、十九畏（丁香+郁金、巴豆+牵牛子）均能准确检测；无禁忌药材对（当归+枸杞子）正确返回无冲突；多味药材场景下仅检测出实际存在的冲突，不产生误报。", "Normal")
add_para(cell, "（5）JSON导入导出正确：导出文件格式合规，重新导入后数据完整一致。", "Normal")
add_para(cell, "（6）数据库重置正确：重置后数据清空，种子数据恢复后数量与原种子数据完全一致。", "Normal")
add_para(cell, "（7）错误处理健全：所有异常路径（重复添加、名称必填校验、药材A=B等）均有正确的提示或返回值。", "Normal")
add_para(cell, "此外，测试使用临时数据库完全隔离，不影响正式数据文件，确保了开发与测试过程的并行开展。", "Normal")
add_para(cell, "结论：系统各项功能均达到设计目标，测试通过。", "Normal")

add_para(cell, "6.4 Bug记录与修改说明", "Heading 2")
add_table_to_cell(cell,
    ["序号", "问题描述", "原因分析", "解决方案", "状态"],
    [
        ["1", "表单数据获取末尾多余换行符", 'tk.Text的get("1.0", "end")在末尾包含自动换行符', '改用get("1.0", "end-1c")去掉末尾字符', "已解决"],
        ["2", "配伍检查输入药材顺序影响匹配结果", "规则存储和输入药材未统一排序", "添加时对herb_a/herb_b做sorted排序存储；检查时对输入也排序比对", "已解决"],
        ["3", "药材表单选择后未清空残留数据", "_on_herb_select未清空旧数据", "在填充前先调用_clear_herb_form()清空所有Text控件", "已解决"],
        ["4", "test.py运行路径问题", "工作目录不一定是项目根目录", "在文件开头使用os.chdir(os.path.dirname(os.path.abspath(__file__)))", "已解决"],
        ["5", "数据库重置后UI列表未刷新", "_reset_database未调用全部refresh方法", "添加对药材列表、规则列表、可选列表、统计、下拉选项的全面刷新", "已解决"],
        ["6", "导入数据后UI未同步更新", "_import_data成功后未刷新列表", "在导入成功后调用全部refresh方法", "已解决"],
    ]
)

# ══════════════════════════════════════════════════════════
#  第七章：用户使用说明 (row 6)
# ══════════════════════════════════════════════════════════

cell = main_table.rows[6].cells[0]
clear_cell(cell)

add_para(cell, "七、用户使用说明", "Heading 1")

add_para(cell, "7.1 系统安装步骤", "Heading 2")
add_para(cell, "（1）安装Python 3.10或更高版本。", "Normal")
add_para(cell, "    下载地址：https://www.python.org/downloads/", "Normal")
add_para(cell, '    安装时请勾选"Add Python to PATH"。', "Normal")
add_para(cell, "（2）下载项目文件。", "Normal")
add_para(cell, "    将项目文件夹（python_keshe/）复制到本地任意目录。", "Normal")
add_para(cell, "（3）无需安装任何第三方依赖。", "Normal")
add_para(cell, "    本项目仅使用Python内置标准库（sqlite3, tkinter, json, os, tempfile），直接可以运行。", "Normal")
add_para(cell, "（4）验证安装。", "Normal")
add_para(cell, '    打开终端（命令提示符或PowerShell），输入：', "Normal")
add_para(cell, "    python --version", "Normal")
add_para(cell, "    确认显示Python 3.10.x或更高版本。", "Normal")

add_para(cell, "7.2 数据库导入方法", "Heading 2")
add_para(cell, "（1）首次运行自动导入：直接运行 python main.py，系统检测到数据库为空时自动导入内置种子数据（25种药材、17条禁忌规则）。", "Normal")
add_para(cell, '（2）从JSON文件导入：启动系统后，点击"数据操作"标签页 → "从JSON文件导入数据"按钮 → 选择.json文件 → 确认导入。', "Normal")
add_para(cell, "JSON文件格式要求：", "Normal")
add_code_para(cell, '{')
add_code_para(cell, '    "herbs": [{ "name": "...", "pinyin": "...", ... }, ...],')
add_code_para(cell, '    "incompatibilities": [{ "herb_a": "...", "herb_b": "...",')
add_code_para(cell, '        "rule_type": "...", "description": "..." }, ...]')
add_code_para(cell, '}')

add_para(cell, "7.3 运行方法", "Heading 2")
add_para(cell, "（1）启动桌面应用：打开终端，进入项目目录，执行：", "Normal")
add_code_para(cell, "python main.py")
add_para(cell, "（2）运行测试套件：", "Normal")
add_code_para(cell, "python test.py")
add_para(cell, "（3）单独使用数据库模块：", "Normal")
add_code_para(cell, "from database import DBManager")
add_code_para(cell, "db = DBManager()  # 连接tcm.db")
add_code_para(cell, "# 执行查询等操作...")
add_code_para(cell, "db.close()")

add_para(cell, "7.4 使用流程说明", "Heading 2")
add_para(cell, "7.4.1 药材查询流程", "Heading 3")
add_para(cell, '步骤1：点击"药材查询"标签页。', "Normal")
add_para(cell, '步骤2：在搜索框中输入关键词（如"甘草"）。', "Normal")
add_para(cell, "步骤3：选择搜索范围（名称/拼音/分类/功效/别名）。", "Normal")
add_para(cell, '步骤4：点击"查询"按钮或按回车键。', "Normal")
add_para(cell, "步骤5：在结果列表中点击药材，底部展开该药材完整信息。", "Normal")

add_para(cell, "7.4.2 药材管理流程", "Heading 3")
add_para(cell, '步骤1：点击"药材管理"标签页。', "Normal")
add_para(cell, '步骤2（新增）：在右侧表单填写药材信息（名称必填），点击"新增药材"。', "Normal")
add_para(cell, '步骤3（修改）：在左侧列表选择药材→修改表单内容→点击"更新所选"。', "Normal")
add_para(cell, '步骤4（删除）：在左侧列表选择药材→点击"删除所选"→确认。', "Normal")

add_para(cell, "7.4.3 配伍禁忌检查流程", "Heading 3")
add_para(cell, '步骤1：点击"配伍禁忌检查"标签页。', "Normal")
add_para(cell, '步骤2：在右侧"可选药材"列表中双击药材名称，将其加入左侧待检查列表。', "Normal")
add_para(cell, "步骤3：确认至少选择2味药材。", "Normal")
add_para(cell, '步骤4：点击"执行配伍禁忌检查"按钮。', "Normal")
add_para(cell, "步骤5：查看底部结果区域——绿色提示表示未发现配伍禁忌；红色提示列出所有冲突药材对及禁忌类型。", "Normal")

add_para(cell, "7.4.4 禁忌规则管理流程", "Heading 3")
add_para(cell, '步骤1：点击"禁忌规则管理"标签页。', "Normal")
add_para(cell, '步骤2（新增）：在右侧表单选择药材A、药材B和禁忌类型，填写说明，点击"新增规则"。', "Normal")
add_para(cell, '步骤3（修改）：在左侧列表选择规则→修改表单→点击"更新规则"。', "Normal")
add_para(cell, '步骤4（删除）：在左侧列表选择规则→点击"删除规则"→确认。', "Normal")

add_para(cell, "7.4.5 数据备份与恢复流程", "Heading 3")
add_para(cell, '步骤1：点击"数据操作"标签页。', "Normal")
add_para(cell, '步骤2（备份）：点击"导出数据到JSON文件"→选择保存路径→确认。', "Normal")
add_para(cell, '步骤3（恢复）：点击"从JSON文件导入数据"→选择.json文件→确认。', "Normal")
add_para(cell, '步骤4（重置）：点击"清空/重置数据库"→确认两次→数据库恢复至初始状态。', "Normal")

# ══════════════════════════════════════════════════════════
#  第八章：总结与体会 (row 7)
# ══════════════════════════════════════════════════════════

cell = main_table.rows[7].cells[0]
clear_cell(cell)

add_para(cell, "八、总结与体会", "Heading 1")

add_para(cell, "8.1 项目完成情况", "Heading 2")
add_para(cell, "本项目按照软件工程标准流程，完成了中医药常用药材查询与配伍禁忌系统的独立设计与实现。项目交付成果如下：", "Normal")
add_table_to_cell(cell,
    ["交付成果", "完成情况", "说明"],
    [
        ["需求分析", "已完成", "明确了三大痛点、五类用户场景、六大功能模块"],
        ["系统设计", "已完成", "数据库设计（2张表，共16个字段）、GUI布局设计"],
        ["系统实现", "已完成", "5个Python源文件，共1513行代码"],
        ["种子数据", "已完成", "25种常用药材 + 17条十八反/十九畏核心禁忌规则"],
        ["测试套件", "已完成", "8大测试场景，17个测试用例，30+条断言，100%通过"],
        ["GUI界面", "已完成", "6个功能标签页，1100×750px桌面应用"],
        ["课程设计报告", "已完成", "含需求分析、系统设计、实现说明、测试分析等"],
    ]
)
add_para(cell, "项目实现了全部预期功能目标，所有测试用例通过。", "Normal")

add_para(cell, "8.2 遇到的主要困难", "Heading 2")
add_para(cell, '（1）tkinter表单数据获取问题：tkinter的Text控件通过get("1.0", "end")获取值时会在末尾附带换行符，导致存储到数据库的数据不干净。解决方案：统一使用get("1.0", "end-1c")在获取时去掉末尾字符。', "Normal")
add_para(cell, '（2）配伍禁忌匹配的顺序问题：用户输入药材的顺序可能与数据库中规则存储的顺序不一致（如输入"甘遂,甘草" vs 存储"甘草,甘遂"）。解决方案：在存储和比对时均对药材名称做sorted排序，确保匹配不受输入顺序影响。', "Normal")
add_para(cell, "（3）UI数据同步问题：数据增删改操作后，多个标签页的数据需要同步刷新。解决方案：在每种数据变更操作完成后，调用对应的_refresh方法，并确保导入/重置等跨模块操作调用所有相关的刷新方法。", "Normal")
add_para(cell, "（4）测试隔离问题：测试运行时不能影响正式数据库。解决方案：使用tempfile在系统临时目录创建独立测试数据库，测试结束后自动删除。", "Normal")
add_para(cell, "（5）数据结构的设计权衡：药材信息字段较多（10个），在设计时需要考虑字段拆分粒度。最终采用单表存储药材全部信息，虽存在少量冗余可能，但简化了查询逻辑，适合本系统的数据规模。", "Normal")

add_para(cell, "8.3 对软件工程流程的理解", "Heading 2")
add_para(cell, "通过本项目的完整实践，我对软件工程的核心流程有了更深刻的认识：", "Normal")
add_para(cell, '（1）需求分析是基石。在编码前投入时间梳理用户需求、设计用例和数据字典，使得后续开发和测试有明确的目标参照，大大减少了后期返工。本项目在需求阶段明确了"配伍禁忌自动检查"这一核心痛点，后续所有设计都围绕它展开。', "Normal")
add_para(cell, "（2）系统设计决定质量。数据库schema的设计直接影响了代码复杂度。本项目herbs表和incompatibilities表的设计经过了字段粒度、约束策略的仔细考量，最终使得CRUD操作简洁可靠，配伍检查算法清晰高效。", "Normal")
add_para(cell, '（3）测试不是事后工作。本项目将测试与开发同步推进，每完成一个模块就编写对应测试用例。测试套件不仅是验收工具，更是开发过程中的"安全网"——每次修改后运行test.py即可快速验证功能未被破坏。', "Normal")
add_para(cell, '（4）迭代优于完美。开发过程中遵循"先跑通核心链路，再逐步完善细节"的策略。例如GUI界面先搭建基本框架和数据绑定，再逐步添加错误提示、确认对话框等用户体验细节。', "Normal")
add_para(cell, "（5）文档与代码同样重要。清晰的项目结构、命名规范和必要的注释，使得即使间隔数周再回看代码也能快速理解。CLAUDE.md文件的维护也让AI辅助开发更加高效。", "Normal")

add_para(cell, "8.4 本课程的收获", "Heading 2")
add_para(cell, '（1）技术层面：熟练掌握了Python sqlite3数据库操作、tkinter GUI编程、JSON数据序列化、模块化代码组织等技能。特别是"零依赖"的技术选型理念——用内置库解决实际问题，体会到Python标准库的强大。', "Normal")
add_para(cell, "（2）工程层面：完整经历了一个软件项目从需求分析→系统设计→编码实现→测试验证→文档撰写的全生命周期，理解了各阶段之间的衔接关系和各自的重要性。", "Normal")
add_para(cell, '（3）领域层面：通过整理中药材数据和配伍规则，深入了解了中医药"十八反""十九畏"的核心内容和临床意义，体会到计算机技术服务于传统医学的价值。', "Normal")
add_para(cell, "（4）方法论层面：学会了如何制定合理的项目进度计划、如何设计可测试的代码结构、如何编写规范的测试用例、如何通过版本管理和AI辅助工具提升开发效率。", "Normal")
add_para(cell, "（5）综合素质层面：独立完成一个完整的软件项目，锻炼了问题分析能力、自主解决问题能力和项目管理能力。遇到困难时学会先分析根因而非盲目尝试，这种思维方式将受益长远。", "Normal")
add_para(cell, '本次课程设计将软件工程理论与具体实践紧密结合，使我对"软件工程综合实践"这门课程有了从理论到实践的完整理解，为今后参与更大规模的软件开发工作打下了坚实基础。', "Normal")

# ══════════════════════════════════════════════════════════
#  第九章：完整代码 (row 8)
# ══════════════════════════════════════════════════════════

cell = main_table.rows[8].cells[0]
clear_cell(cell)

add_para(cell, "九、完整代码", "Heading 1")

source_files = [
    ("main.py", "主程序入口"),
    ("database.py", "数据库模块"),
    ("ui.py", "界面模块（tkinter GUI）"),
    ("data.py", "种子数据模块"),
    ("test.py", "测试套件"),
]

for filename, desc in source_files:
    filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
    with open(filepath, "r", encoding="utf-8") as f:
        code = f.read()

    add_para(cell, f"9.{source_files.index((filename, desc))+1} {filename} — {desc}", "Heading 2")
    add_para(cell, f"文件：{filename}（{len(code.splitlines())}行）", "Normal")
    add_para(cell, "", "Normal")

    for line in code.splitlines():
        # 替换制表符为空格，避免XML问题
        display_line = line.replace("\t", "    ")
        if not display_line:
            display_line = " "
        add_code_para(cell, display_line)

    add_para(cell, "", "Normal")

# ══════════════════════════════════════════════════════════
#  保存
# ══════════════════════════════════════════════════════════

doc.save(DST)
print(f"报告已生成：{DST}")
print("完成！")